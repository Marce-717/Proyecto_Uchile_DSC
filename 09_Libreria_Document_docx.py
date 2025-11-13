# -*- coding: utf-8 -*-
"""
Created on Wed Nov 13 2025

@author: marce
"""
"""
Autor: Marcelo Toro Miranda

"""
"""
FUNDAMENTOS DE PROGRAMACIÓN EN PYTHON
Data Science en Sistemas Naturales
Facultad de Ciencias Agronómicas - Universidad de Chile

Librería python-docx para generación de documentos Word
Ejemplos prácticos para agricultura de precisión y sistemas agropecuarios

"""

# =============================================================================
# 9. LIBRERÍA PYTHON-DOCX - GENERACIÓN DE DOCUMENTOS WORD
# =============================================================================

print("\n\n" + "=" * 60)
print("9. LIBRERÍA PYTHON-DOCX")
print("=" * 60)

# --------------- INSTALACIÓN ---------------

print("\n>>> INSTALACIÓN DE LA LIBRERÍA")
print("-" * 40)
print("""
Para instalar python-docx, ejecutar en la terminal:
    pip install python-docx

Esta librería permite crear y modificar documentos .docx (Word)
de forma programática, ideal para generar reportes automatizados.
""")

# --------------- IMPORTAR LA LIBRERÍA ---------------

print("\n>>> IMPORTAR LA LIBRERÍA")
print("-" * 40)

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    print("✓ Librería python-docx importada correctamente")
except ImportError:
    print("⚠ Error: python-docx no está instalada")
    print("Ejecute: pip install python-docx")
    exit()


# =============================================================================
# EJEMPLO 1: CREAR DOCUMENTO BÁSICO
# =============================================================================

print("\n\n>>> EJEMPLO 1: CREAR DOCUMENTO BÁSICO")
print("-" * 40)

# Crear un nuevo documento
doc = Document()

# Agregar título
doc.add_heading('Reporte de Análisis de Suelo', 0)

# Agregar párrafo
doc.add_paragraph('Este documento contiene los resultados del análisis de suelo.')

# Agregar subtítulo
doc.add_heading('1. Introducción', level=1)

# Agregar párrafo con formato
parrafo = doc.add_paragraph(
    'El análisis de suelo es fundamental para determinar las necesidades '
    'nutricionales del cultivo y optimizar la fertilización.'
)

# Guardar documento
nombre_archivo_1 = 'ejemplo_01_basico.docx'
doc.save(nombre_archivo_1)
print(f"✓ Documento '{nombre_archivo_1}' creado exitosamente")


# =============================================================================
# EJEMPLO 2: DOCUMENTO CON FORMATO AVANZADO
# =============================================================================

print("\n\n>>> EJEMPLO 2: FORMATO AVANZADO")
print("-" * 40)

doc2 = Document()

# Título principal
titulo = doc2.add_heading('Informe Técnico: Análisis de Rendimiento', 0)

# Agregar información del proyecto
doc2.add_heading('Información del Proyecto', level=1)

info_parrafo = doc2.add_paragraph()
info_parrafo.add_run('Predio: ').bold = True
info_parrafo.add_run('Santa Rosa\n')
info_parrafo.add_run('Cultivo: ').bold = True
info_parrafo.add_run('Maíz\n')
info_parrafo.add_run('Temporada: ').bold = True
info_parrafo.add_run('2024-2025\n')

# Agregar viñetas
doc2.add_heading('Objetivos del Análisis', level=1)
doc2.add_paragraph('Evaluar rendimiento por zona de manejo', style='List Bullet')
doc2.add_paragraph('Identificar factores limitantes', style='List Bullet')
doc2.add_paragraph('Generar recomendaciones de mejora', style='List Bullet')

# Agregar párrafo con formato personalizado
doc2.add_heading('Metodología', level=1)
metodologia = doc2.add_paragraph(
    'Se realizó un muestreo sistemático en 15 puntos distribuidos '
    'uniformemente en el predio. Las muestras fueron analizadas en '
    'laboratorio certificado.'
)

# Modificar formato del último párrafo
run = metodologia.runs[0]
run.font.size = Pt(11)
run.font.name = 'Arial'

nombre_archivo_2 = 'ejemplo_02_formato.docx'
doc2.save(nombre_archivo_2)
print(f"✓ Documento '{nombre_archivo_2}' creado con formato avanzado")


# =============================================================================
# EJEMPLO 3: DOCUMENTO CON TABLAS
# =============================================================================

print("\n\n>>> EJEMPLO 3: AGREGAR TABLAS")
print("-" * 40)

doc3 = Document()

doc3.add_heading('Resultados de Análisis de Suelo', 0)
doc3.add_heading('1. Datos de Muestreo', level=1)

# Datos de ejemplo
datos_suelo = [
    {"muestra": "M-01", "ph": 6.5, "n": 25, "p": 18, "k": 180},
    {"muestra": "M-02", "ph": 6.2, "n": 22, "p": 15, "k": 165},
    {"muestra": "M-03", "ph": 6.8, "n": 28, "p": 20, "k": 195},
    {"muestra": "M-04", "ph": 5.9, "n": 19, "p": 12, "k": 150},
    {"muestra": "M-05", "ph": 6.4, "n": 24, "p": 17, "k": 175},
]

# Crear tabla (6 filas x 5 columnas)
# Fila 0 = encabezados, filas 1-5 = datos
tabla = doc3.add_table(rows=len(datos_suelo) + 1, cols=5)
tabla.style = 'Light Grid Accent 1'

# Agregar encabezados
encabezados = ['Muestra', 'pH', 'N (ppm)', 'P (ppm)', 'K (ppm)']
celdas_encabezado = tabla.rows[0].cells
for i, encabezado in enumerate(encabezados):
    celdas_encabezado[i].text = encabezado
    # Hacer los encabezados en negrita
    for paragraph in celdas_encabezado[i].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

# Agregar datos
for i, muestra in enumerate(datos_suelo, start=1):
    fila = tabla.rows[i].cells
    fila[0].text = muestra['muestra']
    fila[1].text = str(muestra['ph'])
    fila[2].text = str(muestra['n'])
    fila[3].text = str(muestra['p'])
    fila[4].text = str(muestra['k'])

# Agregar interpretación
doc3.add_heading('2. Interpretación', level=1)
interpretacion = doc3.add_paragraph(
    'Los resultados muestran valores de pH en el rango óptimo (6.0-7.0) '
    'para la mayoría de las muestras. El contenido de nitrógeno es moderado, '
    'mientras que el fósforo requiere atención en las muestras M-02 y M-04.'
)

nombre_archivo_3 = 'ejemplo_03_tablas.docx'
doc3.save(nombre_archivo_3)
print(f"✓ Documento '{nombre_archivo_3}' creado con tablas")


# =============================================================================
# EJEMPLO 4: REPORTE COMPLETO AUTOMATIZADO
# =============================================================================

print("\n\n>>> EJEMPLO 4: REPORTE COMPLETO AUTOMATIZADO")
print("-" * 40)

def crear_reporte_cultivo(datos_predio, datos_rendimiento):
    """
    Genera un reporte completo de evaluación de cultivo

    Parámetros:
        datos_predio: diccionario con información del predio
        datos_rendimiento: lista de diccionarios con datos por zona

    Retorna:
        nombre del archivo generado
    """
    doc = Document()

    # === PORTADA ===
    doc.add_heading('REPORTE DE EVALUACIÓN DE CULTIVO', 0)

    # Información del predio
    doc.add_paragraph()  # Espacio
    info = doc.add_paragraph()
    info.add_run('Predio: ').bold = True
    info.add_run(f"{datos_predio['nombre']}\n")
    info.add_run('Ubicación: ').bold = True
    info.add_run(f"{datos_predio['ubicacion']}\n")
    info.add_run('Cultivo: ').bold = True
    info.add_run(f"{datos_predio['cultivo']}\n")
    info.add_run('Superficie: ').bold = True
    info.add_run(f"{datos_predio['superficie']} ha\n")
    info.add_run('Fecha de evaluación: ').bold = True
    info.add_run(f"{datos_predio['fecha']}\n")

    # === RESUMEN EJECUTIVO ===
    doc.add_page_break()
    doc.add_heading('1. Resumen Ejecutivo', level=1)

    # Calcular estadísticas
    rendimientos = [z['rendimiento'] for z in datos_rendimiento]
    prom_rend = sum(rendimientos) / len(rendimientos)
    max_rend = max(rendimientos)
    min_rend = min(rendimientos)

    resumen = doc.add_paragraph(
        f"El cultivo de {datos_predio['cultivo']} en el predio "
        f"{datos_predio['nombre']} presentó un rendimiento promedio de "
        f"{prom_rend:.2f} ton/ha, con valores que oscilaron entre "
        f"{min_rend:.2f} y {max_rend:.2f} ton/ha."
    )

    # === RESULTADOS POR ZONA ===
    doc.add_heading('2. Resultados por Zona de Manejo', level=1)

    # Crear tabla de rendimientos
    tabla_rend = doc.add_table(rows=len(datos_rendimiento) + 1, cols=4)
    tabla_rend.style = 'Light Grid Accent 1'

    # Encabezados
    headers = ['Zona', 'Superficie (ha)', 'Rendimiento (ton/ha)', 'Producción (ton)']
    for i, header in enumerate(headers):
        cell = tabla_rend.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Datos
    total_superficie = 0
    total_produccion = 0

    for i, zona in enumerate(datos_rendimiento, start=1):
        row = tabla_rend.rows[i].cells
        superficie = zona['superficie']
        rendimiento = zona['rendimiento']
        produccion = superficie * rendimiento

        row[0].text = zona['zona']
        row[1].text = f"{superficie:.1f}"
        row[2].text = f"{rendimiento:.2f}"
        row[3].text = f"{produccion:.2f}"

        total_superficie += superficie
        total_produccion += produccion

    # === ANÁLISIS Y RECOMENDACIONES ===
    doc.add_heading('3. Análisis y Recomendaciones', level=1)

    doc.add_heading('3.1 Variabilidad Espacial', level=2)
    doc.add_paragraph(
        f"Se observa una variabilidad de {(max_rend - min_rend):.2f} ton/ha "
        f"entre las zonas de mayor y menor rendimiento, lo que representa "
        f"un {((max_rend - min_rend) / prom_rend * 100):.1f}% respecto al "
        f"promedio."
    )

    doc.add_heading('3.2 Recomendaciones', level=2)
    doc.add_paragraph(
        'Realizar análisis de suelo específicos en zonas de bajo rendimiento',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Evaluar sistema de riego para optimizar distribución de agua',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Implementar fertilización variable según mapas de rendimiento',
        style='List Bullet'
    )

    # === PIE DE PÁGINA ===
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run('_' * 60 + '\n')
    footer.add_run('Generado automáticamente con Python - python-docx\n')
    footer.add_run('Facultad de Ciencias Agronómicas - Universidad de Chile')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Guardar documento
    nombre_archivo = f"reporte_{datos_predio['nombre'].replace(' ', '_')}.docx"
    doc.save(nombre_archivo)

    return nombre_archivo


# Datos de ejemplo para el reporte
datos_predio_ejemplo = {
    'nombre': 'Santa Elena',
    'ubicacion': 'Región Metropolitana',
    'cultivo': 'Trigo',
    'superficie': 45.5,
    'fecha': '13 de Noviembre 2025'
}

datos_rendimiento_ejemplo = [
    {'zona': 'Zona A', 'superficie': 12.5, 'rendimiento': 6.8},
    {'zona': 'Zona B', 'superficie': 15.3, 'rendimiento': 7.2},
    {'zona': 'Zona C', 'superficie': 8.7, 'rendimiento': 5.9},
    {'zona': 'Zona D', 'superficie': 9.0, 'rendimiento': 6.5},
]

# Generar reporte
archivo_generado = crear_reporte_cultivo(datos_predio_ejemplo, datos_rendimiento_ejemplo)
print(f"✓ Reporte completo generado: '{archivo_generado}'")


# =============================================================================
# EJEMPLO 5: AGREGAR ESTILOS Y FORMATO PERSONALIZADO
# =============================================================================

print("\n\n>>> EJEMPLO 5: ESTILOS Y FORMATO PERSONALIZADO")
print("-" * 40)

doc5 = Document()

# Título con formato personalizado
titulo = doc5.add_heading('Protocolo de Muestreo de Suelo', 0)

# Párrafo con diferentes estilos en el mismo texto
doc5.add_heading('1. Materiales Necesarios', level=1)

materiales = doc5.add_paragraph()
materiales.add_run('Importante: ').bold = True
materiales.add_run('Utilizar ')
materiales.add_run('únicamente').italic = True
materiales.add_run(' herramientas limpias y desinfectadas.')

# Lista numerada
doc5.add_heading('2. Procedimiento', level=1)
doc5.add_paragraph('Identificar zonas de muestreo homogéneas', style='List Number')
doc5.add_paragraph('Tomar 10-15 submuestras por zona', style='List Number')
doc5.add_paragraph('Mezclar submuestras y tomar muestra compuesta', style='List Number')
doc5.add_paragraph('Enviar al laboratorio en plazo de 24 horas', style='List Number')

# Texto con color personalizado
doc5.add_heading('3. Consideraciones Importantes', level=1)
advertencia = doc5.add_paragraph()
run_advertencia = advertencia.add_run(
    '⚠ IMPORTANTE: Las muestras deben ser tomadas a la misma profundidad '
    'en todas las zonas (0-20 cm para cultivos anuales).'
)
run_advertencia.font.color.rgb = RGBColor(255, 0, 0)  # Rojo
run_advertencia.font.size = Pt(12)

nombre_archivo_5 = 'ejemplo_05_estilos.docx'
doc5.save(nombre_archivo_5)
print(f"✓ Documento '{nombre_archivo_5}' creado con estilos personalizados")


# =============================================================================
# RESUMEN DE FUNCIONALIDADES
# =============================================================================

print("\n\n>>> RESUMEN DE FUNCIONALIDADES PRINCIPALES")
print("=" * 60)
print("""
FUNCIONES BÁSICAS:
• Document()                      - Crear nuevo documento
• doc.add_heading()               - Agregar títulos y subtítulos
• doc.add_paragraph()             - Agregar párrafos
• doc.add_page_break()            - Insertar salto de página
• doc.save()                      - Guardar documento

FORMATO DE TEXTO:
• paragraph.add_run()             - Agregar texto con formato
• run.bold = True                 - Texto en negrita
• run.italic = True               - Texto en cursiva
• run.font.size = Pt(12)          - Cambiar tamaño de fuente
• run.font.color.rgb              - Cambiar color de texto

TABLAS:
• doc.add_table(rows, cols)       - Crear tabla
• table.style                     - Aplicar estilo a tabla
• table.rows[i].cells[j].text     - Acceder a celdas

ALINEACIÓN:
• paragraph.alignment             - Alinear texto
• WD_ALIGN_PARAGRAPH.CENTER       - Centrar
• WD_ALIGN_PARAGRAPH.RIGHT        - Alinear a la derecha

LISTAS:
• style='List Bullet'             - Lista con viñetas
• style='List Number'             - Lista numerada

APLICACIONES EN AGRICULTURA:
✓ Reportes de análisis de suelo
✓ Informes de rendimiento de cultivos
✓ Protocolos y procedimientos técnicos
✓ Documentación de ensayos
✓ Reportes de fertilización
✓ Evaluaciones de impacto ambiental
""")

print("\n" + "=" * 60)
print("Documentos de ejemplo creados exitosamente")
print("=" * 60)
print("\nArchivos generados:")
print(f"  • {nombre_archivo_1}")
print(f"  • {nombre_archivo_2}")
print(f"  • {nombre_archivo_3}")
print(f"  • {archivo_generado}")
print(f"  • {nombre_archivo_5}")
print("\n✓ Tutorial completado")
